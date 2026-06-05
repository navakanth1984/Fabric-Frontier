import sys
from docx import Document
from pathlib import Path

def convert_to_docx(input_file, output_file):
    if not Path(input_file).exists():
        print(f"Error: {input_file} not found.")
        return

    doc = Document()
    # Set default style for screenplays
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'

    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.isupper() and (line.startswith('INT.') or line.startswith('EXT.')):
            # Scene heading
            p = doc.add_paragraph()
            p.add_run(line).bold = True
        elif line.isupper() and ':' not in line:
            # Character name - centered-ish in screenplay style
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = 200000 # Rough indent
            p.add_run(line).bold = True
        else:
            doc.add_paragraph(line)

    doc.save(output_file)
    print(f"Successfully converted {input_file} to {output_file}")

if __name__ == '__main__':
    convert_to_docx('factory_output.md', 'DAAVA_Factory_Scene_V1.docx')
