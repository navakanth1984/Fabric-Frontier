import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_screenplay(input_md, output_docx):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
    style = doc.styles['Normal']
    style.font.name = 'Courier New'
    style.font.size = Pt(12)
    
    if not os.path.exists(input_md):
        print(f"Error: {input_md} not found.")
        return

    with open(input_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    last_was_character = False
    last_was_paren = False
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line: 
            last_was_character = False; last_was_paren = False; continue
        
        if line.startswith('INT.') or line.startswith('EXT.'):
            p = doc.add_paragraph()
            p.add_run(line.upper()).bold = True
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            last_was_character = False
        elif line.endswith(':') and line.isupper() and len(line.split()) <= 3:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(12)
            last_was_character = False
        elif line.isupper() and not any(c in line for c in ['.', ',', '!', '?', '-']) and len(line.split()) <= 4:
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(2.2)
            p.paragraph_format.space_before = Pt(12)
            last_was_character = True; last_was_paren = False
        elif line.startswith('(') and line.endswith(')'):
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(1.6)
            last_was_character = False; last_was_paren = True
        elif last_was_character or last_was_paren:
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(1.0)
            p.paragraph_format.right_indent = Inches(1.0)
            last_was_character = False; last_was_paren = False
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_before = Pt(12)
            last_was_character = False
            
    doc.save(output_docx)
    print(f'Successfully formatted {input_md} to {output_docx}')

if __name__ == '__main__':
    in_file = sys.argv[1] if len(sys.argv) > 1 else 'factory_output.md'
    out_file = sys.argv[2] if len(sys.argv) > 2 else 'DAAVA_Screenplay_Formatted.docx'
    format_screenplay(in_file, out_file)
